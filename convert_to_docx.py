#!/usr/bin/env python3
"""
Convert all Sibelius curriculum Markdown lessons to professional DOCX (Word) files.
Output: 8 files organized by semester, with Arabic font support and formatting.
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import glob

BASE_DIR = "/opt/data/sibelius-curriculum"
OUTPUT_DIR = os.path.join(BASE_DIR, "docx_output")

SEMESTER1_DIR = os.path.join(BASE_DIR, "01-Semester1", "Lessons")
SEMESTER2_DIR = os.path.join(BASE_DIR, "02-Semester2", "Lessons")

os.makedirs(OUTPUT_DIR, exist_ok=True)

def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_page_number(doc):
    """Add page number to footer."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = paragraph.add_run()
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._r.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)

def parse_markdown_simple(md_text):
    """
    Parse Markdown into structured elements for DOCX conversion.
    Returns list of (type, content) tuples.
    """
    elements = []
    lines = md_text.split('\n')
    
    i = 0
    in_code_block = False
    code_buffer = []
    in_table = False
    table_buffer = []
    
    while i < len(lines):
        line = lines[i]
        
        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                elements.append(('code', '\n'.join(code_buffer)))
                code_buffer = []
                in_code_block = False
                i += 1
                continue
            else:
                in_code_block = True
                i += 1
                continue
        
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue
        
        # Tables
        if '|' in line and line.strip().startswith('|'):
            in_table = True
            table_buffer.append(line)
            i += 1
            continue
        else:
            if in_table and len(table_buffer) >= 2:  # header + separator + data
                elements.append(('table', '\n'.join(table_buffer)))
                table_buffer = []
                in_table = False
            elif in_table:
                table_buffer = []
                in_table = False
        
        # Headers
        if line.startswith('# '):
            elements.append(('h1', line[2:].strip()))
        elif line.startswith('## '):
            elements.append(('h2', line[2:].strip()))
        elif line.startswith('### '):
            elements.append(('h3', line[2:].strip()))
        elif line.startswith('#### '):
            elements.append(('h4', line[2:].strip()))
        # Bullet lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            elements.append(('bullet', line.strip()[2:].strip()))
        # Numbered lists
        elif re.match(r'^\d+[\.\)] ', line.strip()):
            elements.append(('numbered', re.sub(r'^\d+[\.\)] ', '', line.strip())))
        # Horizontal rule
        elif line.strip() in ('---', '___', '***'):
            elements.append(('hr', ''))
        # Blank line
        elif line.strip() == '':
            pass
        # Regular paragraph
        else:
            elements.append(('para', line.strip()))
        
        i += 1
    
    # Flush remaining buffers
    if in_code_block and code_buffer:
        elements.append(('code', '\n'.join(code_buffer)))
    if in_table and len(table_buffer) >= 2:
        elements.append(('table', '\n'.join(table_buffer)))
    
    return elements

def parse_table_from_md(table_text):
    """Convert markdown table text to rows of cells."""
    lines = [l.strip() for l in table_text.split('\n') if l.strip()]
    if len(lines) < 2:
        return []
    
    # Skip separator line (|---|---|)
    data_lines = [line for line in lines if not re.match(r'^[\|\s:\-]+$', line)]
    
    rows = []
    for line in data_lines:
        cells = [c.strip() for c in line.split('|') if c.strip()]
        if cells:
            rows.append(cells)
    
    return rows

def add_formatted_para(doc, text, style='para', bold=False, italic=False,
                       font_size=11, color=None, alignment=None, space_after=6, 
                       font_name='Calibri', level=0):
    """Add a formatted paragraph to the document."""
    # Handle inline markdown in text
    p = doc.add_paragraph()
    p.space_after = Pt(space_after)
    p.space_before = Pt(0)
    
    if alignment:
        p.alignment = alignment
    
    if style == 'h1':
        p.style = doc.styles['Heading 1']
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        p.space_before = Pt(24)
        p.space_after = Pt(12)
        return p
    
    if style == 'h2':
        p.style = doc.styles['Heading 2']
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x6B)
        p.space_before = Pt(18)
        p.space_after = Pt(8)
        return p
    
    if style == 'h3':
        p.style = doc.styles['Heading 3']
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x34, 0x49, 0x7E)
        p.space_before = Pt(12)
        p.space_after = Pt(6)
        return p
    
    # For regular paragraphs, use the font settings
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run.bold = bold
    run.italic = italic
    
    if color:
        run.font.color.rgb = RGBColor(*color)
    
    return p

def add_mcq_table(doc, questions):
    """Add MCQ section as a formatted table."""
    if not questions:
        return
    
    # Group questions in pairs (2 per row)
    table = doc.add_table(rows=(len(questions) + 1) // 2 + 1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '2C3E6B')
        borders.append(border)
    tblPr.append(borders)
    
    # Header
    header_cells = table.rows[0].cells
    for i in range(2):
        for j, txt in enumerate(['السؤال', 'الإجابة']):
            cell = header_cells[i * 2 + j]
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, '2C3E6B')
    
    # Content
    for idx, q in enumerate(questions):
        row_idx = idx // 2 + 1
        col_start = (idx % 2) * 2
        q_cell = table.rows[row_idx].cells[col_start]
        a_cell = table.rows[row_idx].cells[col_start + 1]
        
        # Split question number from text
        q_text = q
        if '?' in q:
            parts = q.split('?', 1)
            q_text = parts[0] + '؟'
        
        p1 = q_cell.paragraphs[0]
        run1 = p1.add_run(q_text)
        run1.font.size = Pt(9)
        run1.font.name = 'Calibri'
        
        # Find MCQ options and answer
        a_text = '—'
        for opt in q.split('\n'):
            opt = opt.strip()
            if opt.startswith('**') and ('✔' in opt or '✓' in opt):
                a_text = opt.strip('*').replace('✔', '').replace('✓', '').replace('(', '').replace(')', '').strip()
            elif '✔' in opt or '✓' in opt:
                a_text = opt.replace('✔', '').replace('✓', '').strip()
        
        p2 = a_cell.paragraphs[0]
        run2 = p2.add_run(a_text)
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x1B, 0x7A, 0x2B)
        run2.font.bold = True
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Alternate row shading
        if idx % 2 == 0:
            set_cell_shading(q_cell, 'F0F4FF')
            set_cell_shading(a_cell, 'F0F4FF')

def markdown_to_docx(md_text, doc):
    """Convert parsed markdown elements to DOCX formatting."""
    elements = parse_markdown_simple(md_text)
    
    mcq_buffer = []
    in_mcq = False
    in_homework = False
    homework_buffer = []
    
    for elem_type, content in elements:
        # Detect MCQ section
        if elem_type == 'h2' and ('أسئلة' in content or 'Questions' in content or 'MCQ' in content):
            in_mcq = True
            add_formatted_para(doc, content, style='h2')
            continue
        
        # Detect Homework section
        if elem_type == 'h2' and ('واجب' in content or 'Homework' in content or 'تمرين' in content):
            in_mcq = False
            in_homework = True
            add_formatted_para(doc, content, style='h2')
            continue
        
        # End of homework
        if in_homework and elem_type in ('h1', 'h2', 'h3'):
            in_homework = False
        
        if in_mcq:
            if elem_type == 'para' and content:
                # Check if it's a question line
                if re.match(r'^\d+[\.\)]', content):
                    mcq_buffer.append(content)
                elif 'a)' in content or 'b)' in content or 'c)' in content or 'd)' in content:
                    if mcq_buffer:
                        mcq_buffer[-1] = mcq_buffer[-1] + '\n' + content
                elif content.startswith('**') and ('✔' in content or '✓' in content):
                    if mcq_buffer:
                        mcq_buffer[-1] = mcq_buffer[-1] + '\n' + content
            elif elem_type == 'bullet' and content:
                mcq_buffer.append(content)
            elif elem_type == 'numbered' and content:
                mcq_buffer.append(content)
            elif elem_type == 'h2' or elem_type == 'h3':
                # MCQ section ended
                if mcq_buffer:
                    add_mcq_table(doc, mcq_buffer)
                    mcq_buffer = []
                in_mcq = False
                add_formatted_para(doc, content, style='h2')
            continue
        
        # Regular elements
        if elem_type == 'h1':
            add_formatted_para(doc, content, style='h1')
        elif elem_type == 'h2':
            add_formatted_para(doc, content, style='h2')
        elif elem_type == 'h3':
            add_formatted_para(doc, content, style='h3')
        elif elem_type == 'h4':
            add_formatted_para(doc, content, style='h3')  # Use h3 style
        elif elem_type == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(content)
            run.font.size = Pt(11)
        elif elem_type == 'numbered':
            p = doc.add_paragraph(style='List Number')
            run = p.add_run(content)
            run.font.size = Pt(11)
        elif elem_type == 'para':
            if content:
                is_bold = content.startswith('**') and content.endswith('**')
                clean = content.strip('*')
                add_formatted_para(doc, clean, bold=is_bold, font_size=11)
        elif elem_type == 'code':
            p = doc.add_paragraph()
            p.space_before = Pt(6)
            p.space_after = Pt(6)
            run = p.add_run(content)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            # Add grey background effect
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F5F5F5')
            shading.set(qn('w:val'), 'clear')
            p._p.get_or_add_pPr().append(shading)
        elif elem_type == 'table':
            rows = parse_table_from_md(content)
            if rows:
                num_cols = max(len(r) for r in rows)
                t = doc.add_table(rows=len(rows), cols=num_cols)
                t.style = 'Light Grid Accent 1'
                for ri, row_data in enumerate(rows):
                    for ci, cell_text in enumerate(row_data):
                        if ci < num_cols:
                            cell = t.rows[ri].cells[ci]
                            cell.text = ''
                            p = cell.paragraphs[0]
                            run = p.add_run(cell_text)
                            run.font.size = Pt(9)
                            if ri == 0:  # header row
                                run.font.bold = True
    
    # Flush remaining buffers
    if mcq_buffer:
        add_mcq_table(doc, mcq_buffer)

def process_lesson_file(filepath, doc):
    """Process a single lesson markdown file into an existing document."""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Add page break before each lesson (except first)
    doc.add_page_break()
    
    markdown_to_docx(content, doc)

def create_semester_doc(lesson_files, output_name, semester_title):
    """Create a DOCX file for a group of lessons."""
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # Default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Heading styles
    for i in range(1, 4):
        h_style = doc.styles[f'Heading {i}']
        h_font = h_style.font
        h_font.name = 'Calibri'
    
    # Page numbers
    add_page_number(doc)
    
    # ── Cover Page ──
    for _ in range(6):
        doc.add_paragraph()
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(semester_title)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    
    # Subtitle
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("منهج تعليم Sibelius الشامل")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x6B)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Comprehensive Sibelius Curriculum")
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{len(lesson_files)} دروس")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("إعداد: د. مصطفى كمال")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x34, 0x49, 0x7E)
    
    doc.add_page_break()
    
    # ── Table of Contents ──
    add_formatted_para(doc, 'فهرس المحتويات — Table of Contents', style='h1')
    doc.add_paragraph()
    
    for fpath in lesson_files:
        fname = os.path.basename(fpath)
        lesson_num = re.search(r'(\d+)', fname)
        num = lesson_num.group(1) if lesson_num else '?'
        
        p = doc.add_paragraph()
        run = p.add_run(f"الدرس {num} – ")
        run.font.size = Pt(11)
        run.font.bold = True
        
        # Get first heading of the file
        with open(fpath, 'r', encoding='utf-8') as f:
            first_line = f.readline().strip('# \n')
            if not first_line.startswith('#'):
                f.seek(0)
                for line in f:
                    if line.startswith('# '):
                        first_line = line.strip('# \n')
                        break
        
        run2 = p.add_run(first_line)
        run2.font.size = Pt(11)
        run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_page_break()
    
    # ── Lessons ──
    add_formatted_para(doc, f'{semester_title}', style='h1')
    doc.add_paragraph()
    
    for i, fpath in enumerate(lesson_files):
        if i > 0:
            doc.add_page_break()
        with open(fpath, 'r', encoding='utf-8') as f:
            content = f.read()
        markdown_to_docx(content, doc)
    
    # Save
    output_path = os.path.join(OUTPUT_DIR, output_name)
    doc.save(output_path)
    print(f"✅ Saved: {output_path}")
    return output_path


def main():
    # Collect lesson files sorted by number
    s1_files = sorted(glob.glob(os.path.join(SEMESTER1_DIR, "*.md")))
    s2_files = sorted(glob.glob(os.path.join(SEMESTER2_DIR, "*.md")))
    
    print(f"Semester 1: {len(s1_files)} lessons")
    print(f"Semester 2: {len(s2_files)} lessons")
    
    # Split into 4 files per semester (8 lessons each)
    s1_batches = [
        s1_files[0:8],
        s1_files[8:16],
        s1_files[16:24],
        s1_files[24:32],
    ]
    
    s2_batches = [
        s2_files[0:8],
        s2_files[8:16],
        s2_files[16:24],
        s2_files[24:32],
    ]
    
    semester1_titles = [
        "Sibelius Semester 1 — Part 1 (Lessons 1–8): Foundations",
        "Sibelius Semester 1 — Part 2 (Lessons 9–16): Core Skills",
        "Sibelius Semester 1 — Part 3 (Lessons 17–24): Notation & Formatting",
        "Sibelius Semester 1 — Part 4 (Lessons 25–32): Advanced Semester 1",
    ]
    
    semester2_titles = [
        "Sibelius Semester 2 — Part 1 (Lessons 33–40): Intermediate Techniques",
        "Sibelius Semester 2 — Part 2 (Lessons 41–48): Advanced Projects",
        "Sibelius Semester 2 — Part 3 (Lessons 49–56): Professional Workflows",
        "Sibelius Semester 2 — Part 4 (Lessons 57–64): Final Projects & Mastery",
    ]
    
    outputs = []
    
    # Process Semester 1
    print("\n📘 Processing Semester 1...")
    for i, batch in enumerate(s1_batches):
        output_name = f"Sibelius-Semester1-Part{i+1}-Lessons-{i*8+1}-{(i+1)*8}.docx"
        title = semester1_titles[i]
        try:
            path = create_semester_doc(batch, output_name, title)
            outputs.append(path)
        except Exception as e:
            print(f"❌ Error on batch {i+1}: {e}")
    
    # Process Semester 2
    print("\n📗 Processing Semester 2...")
    for i, batch in enumerate(s2_batches):
        output_num_start = i * 8 + 33
        output_num_end = (i + 1) * 8 + 32
        output_name = f"Sibelius-Semester2-Part{i+1}-Lessons-{output_num_start}-{output_num_end}.docx"
        title = semester2_titles[i]
        try:
            path = create_semester_doc(batch, output_name, title)
            outputs.append(path)
        except Exception as e:
            print(f"❌ Error on batch {i+1}: {e}")
    
    # Print summary
    print("\n" + "="*60)
    print("📋 CONVERSION SUMMARY")
    print("="*60)
    total_size = 0
    for path in outputs:
        size_kb = os.path.getsize(path) / 1024
        total_size += size_kb
        print(f"  📄 {os.path.basename(path)} — {size_kb:.0f} KB")
    print(f"\n  📦 Total: {len(outputs)} files, {total_size:.0f} KB")
    print(f"  📂 Location: {OUTPUT_DIR}/")

if __name__ == "__main__":
    main()
