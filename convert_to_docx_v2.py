#!/usr/bin/env python3
"""
Professional Sibelius Curriculum → DOCX Converter v2
Arabic RTL support, proper headings, no markdown symbols, centered titles.
"""

import os, re, glob
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

BASE_DIR = "/opt/data/sibelius-curriculum"
OUTPUT_DIR = os.path.join(BASE_DIR, "docx_output_v2")

SEMESTER1_DIR = os.path.join(BASE_DIR, "01-Semester1", "Lessons")
SEMESTER2_DIR = os.path.join(BASE_DIR, "02-Semester2", "Lessons")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ─── helpers ───

def set_rtl(paragraph):
    """Set paragraph to right-to-left."""
    pPr = paragraph._p.get_or_add_pPr()
    # Remove any existing bidi
    for old in pPr.findall(qn('w:bidi')):
        pPr.remove(old)
    rtl = OxmlElement('w:bidi')
    rtl.set(qn('w:val'), '1')
    pPr.append(rtl)
    # Set text direction
    textDir = OxmlElement('w:textDirection')
    textDir.set(qn('w:val'), 'lrTb')  # right-to-left text flow
    pPr.append(textDir)

def make_run_rtl(run):
    """Force a run to be RTL."""
    rPr = run._r.get_or_add_rPr()
    # Remove existing rtl
    for old in rPr.findall(qn('w:rtl')):
        rPr.remove(old)
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)
    # Also set complex script font for Arabic
    csFont = OxmlElement('w:rFonts')
    csFont.set(qn('w:cs'), 'Calibri')
    rPr.append(csFont)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(run.font.size.pt * 2)) if run.font.size else '22')
    rPr.append(szCs)

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_page_number(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)

def clean_inline(text):
    """Remove **, __, ###, --- markdown symbols from text, and fix bracket direction."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    # Remove ### headings that slipped through as inline text
    text = re.sub(r'^###+ ', '', text, flags=re.MULTILINE)
    # Remove --- markers
    text = re.sub(r'^---+$', '', text, flags=re.MULTILINE)
    # Fix bracket direction: wrap parentheses in Unicode direction marks
    # This prevents Word from flipping ( and ) in RTL context
    LRM = '\u200f'
    text = re.sub(r'\(([^)]+)\)', LRM + r'(\1)' + LRM, text)
    return text

def add_formatted_run(p, text, bold=False, italic=False, size=11, color=None, font='Calibri'):
    """Add a run with formatting to a paragraph."""
    text = clean_inline(text)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    make_run_rtl(run)
    return run

def new_para(doc, text='', style='para', bold=False, size=11, color=None, 
             alignment=None, space_before=0, space_after=6, rtl=True):
    """Create a new paragraph with proper formatting."""
    p = doc.add_paragraph()
    p.space_before = Pt(space_before)
    p.space_after = Pt(space_after)
    if rtl:
        set_rtl(p)
    if alignment:
        p.alignment = alignment
    if text:
        add_formatted_run(p, text, bold=bold, size=size, color=color)
    return p

def add_heading(doc, text, level=1, rtl=True):
    """Add a clean heading — no style corruption."""
    sizes = {1: 22, 2: 18, 3: 14, 4: 12}
    colors = {1: (0x1A, 0x1A, 0x2E), 2: (0x2C, 0x3E, 0x6B), 3: (0x34, 0x49, 0x7E), 4: (0x44, 0x55, 0x88)}
    spaces = {1: 28, 2: 20, 3: 14, 4: 10}
    
    clean = clean_inline(text.strip().rstrip(':'))
    
    p = doc.add_paragraph()
    p.space_before = Pt(spaces[level])
    p.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level <= 2 else WD_ALIGN_PARAGRAPH.RIGHT
    if rtl:
        set_rtl(p)
    
    run = p.add_run(clean)
    run.font.size = Pt(sizes[level])
    run.font.name = 'Calibri'
    run.bold = True
    run.font.color.rgb = RGBColor(*colors[level])
    make_run_rtl(run)
    return p

# ─── markdown parser ───

def parse_md_sections(md_text):
    """Parse markdown into structured blocks."""
    blocks = []
    lines = md_text.split('\n')
    
    i = 0
    in_code = False
    code_buf = []
    
    while i < len(lines):
        line = lines[i]
        
        # Code fences
        if line.strip().startswith('```'):
            if in_code:
                blocks.append(('code', '\n'.join(code_buf)))
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue
        
        # Headings
        if line.startswith('# '):
            blocks.append(('h1', line[2:]))
        elif line.startswith('## '):
            blocks.append(('h2', line[3:]))
        elif line.startswith('### '):
            blocks.append(('h3', line[4:]))
        elif line.startswith('#### '):
            blocks.append(('subheading', line[5:]))
        elif line.startswith('##### '):
            blocks.append(('subheading', line[6:]))
        elif line.strip().startswith('---'):
            blocks.append(('hr', ''))
            i += 1
            continue
        elif line.strip().startswith('|'):
            # Collect table
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                tbl_lines.append(lines[i])
                i += 1
            blocks.append(('table', '\n'.join(tbl_lines)))
            i -= 1  # compensate for loop increment
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            blocks.append(('bullet', line.strip()[2:]))
        elif re.match(r'^\d+[\.\)] ', line.strip()):
            blocks.append(('numbered', re.sub(r'^\d+[\.\)] ', '', line.strip())))
        elif line.strip() == '':
            pass  # skip blank lines
        else:
            # Regular paragraph — might be part of multi-line
            para_lines = [line]
            while i+1 < len(lines) and lines[i+1].strip() and not lines[i+1].strip().startswith('#') and not lines[i+1].strip().startswith('|') and not lines[i+1].strip().startswith('---'):
                i += 1
                para_lines.append(lines[i])
            blocks.append(('para', '\n'.join(para_lines)))
        
        i += 1
    
    # Flush code buffer
    if in_code and code_buf:
        blocks.append(('code', '\n'.join(code_buf)))
    
    return blocks

def parse_table_rows(table_text):
    """Parse markdown table into 2D list."""
    lines = [l.strip() for l in table_text.split('\n') if l.strip()]
    if len(lines) < 2:
        return []
    # Skip separator
    data = [l for l in lines if not re.match(r'^[\|\s:\-]+$', l)]
    rows = []
    for line in data:
        cells = [c.strip() for c in line.split('|') if c.strip()]
        if cells:
            rows.append(cells)
    return rows

def table_to_docx(doc, rows):
    """Render a markdown table into DOCX table."""
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=num_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Borders
    tbl_pr = tbl._tbl.tblPr if tbl._tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for b in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{b}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '2C3E6B')
        borders.append(el)
    tbl_pr.append(borders)
    
    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            if ci >= num_cols:
                break
            cell = tbl.rows[ri].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_rtl(p)
            run = p.add_run(clean_inline(cell_text))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            make_run_rtl(run)
            if ri == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_cell_shading(cell, '2C3E6B')
            elif ri % 2 == 0:
                set_cell_shading(cell, 'F0F4FF')

def process_lesson(doc, md_text):
    """Convert a full lesson markdown to DOCX elements."""
    blocks = parse_md_sections(md_text)
    
    mcq_rows = []
    in_mcq = False
    mcq_section = False
    
    for btype, content in blocks:
        # Detect MCQ / Homework sections
        if btype == 'h2' and ('أسئلة' in content or 'Questions' in content or 'MCQ' in content or 'Quiz' in content):
            add_heading(doc, content, 2)
            mcq_section = True
            in_mcq = True
            mcq_rows = []
            continue
        
        if btype == 'h2' and ('واجب' in content or 'Homework' in content or 'تمرين منزلي' in content):
            add_heading(doc, content, 2)
            mcq_section = False
            in_mcq = False
            if mcq_rows:
                render_mcq_table(doc, mcq_rows)
                mcq_rows = []
            continue
        
        if btype == 'h2' and mcq_section and not in_mcq:
            if mcq_rows:
                render_mcq_table(doc, mcq_rows)
                mcq_rows = []
            mcq_section = False
            add_heading(doc, content, 2)
            continue
        
        if in_mcq and btype in ('para', 'bullet', 'numbered'):
            if content.strip():
                mcq_rows.append(content.strip())
            continue
        
        # End MCQ section at next heading
        if in_mcq and btype.startswith('h'):
            if mcq_rows:
                render_mcq_table(doc, mcq_rows)
                mcq_rows = []
            in_mcq = False
            add_heading(doc, content, int(btype[1]))
            continue
        
        # Flush MCQ if we're out of MCQ mode
        if mcq_rows and not in_mcq:
            render_mcq_table(doc, mcq_rows)
            mcq_rows = []
        
        # Normal blocks
        if btype == 'h1':
            add_heading(doc, content, 1)
        elif btype == 'h2':
            add_heading(doc, content, 2)
        elif btype == 'h3':
            add_heading(doc, content, 3)
        elif btype == 'h4':
            add_heading(doc, content, 4)
        elif btype == 'subheading':
            p = doc.add_paragraph()
            set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(clean_inline(content))
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
            run.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            make_run_rtl(run)
        elif btype == 'bullet':
            p = doc.add_paragraph()
            set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run('• ' + clean_inline(content))
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
            make_run_rtl(run)
        elif btype == 'numbered':
            p = doc.add_paragraph()
            set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(clean_inline(content))
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
            make_run_rtl(run)
        elif btype == 'para':
            if content.strip():
                p = doc.add_paragraph()
                set_rtl(p)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                is_bold = content.strip().startswith('**') and content.strip().endswith('**')
                run = p.add_run(clean_inline(content))
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
                run.bold = is_bold
                make_run_rtl(run)
        elif btype == 'table':
            rows = parse_table_rows(content)
            if rows:
                table_to_docx(doc, rows)
        elif btype == 'code':
            p = doc.add_paragraph()
            set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(clean_inline(content))
            run.font.size = Pt(9)
            run.font.name = 'Courier New'
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            make_run_rtl(run)
            # Light grey bg
            from docx.oxml import OxmlElement
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F5F5F5')
            shading.set(qn('w:val'), 'clear')
            p._p.get_or_add_pPr().append(shading)
    
    # Flush remaining MCQ
    if mcq_rows:
        render_mcq_table(doc, mcq_rows)


def render_mcq_table(doc, rows):
    """Render MCQ questions as a 2-column table."""
    if not rows:
        return
    
    questions = []
    current = None
    
    for row in rows:
        if re.match(r'^\d+[\.\)]', row):
            if current:
                questions.append(current)
            current = row
        elif current and ('a)' in row or 'b)' in row or 'c)' in row or 'd)' in row):
            current += '\n' + row
    
    if current:
        questions.append(current)
    
    if not questions:
        return
    
    # Table: Question | Correct Answer
    tbl = doc.add_table(rows=len(questions) + 1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Borders
    tbl_pr = tbl._tbl.tblPr if tbl._tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for b in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{b}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '2C3E6B')
        borders.append(el)
    tbl_pr.append(borders)
    
    # Header
    for hi, hdr in enumerate(['السؤال', 'الإجابة الصحيحة']):
        cell = tbl.rows[0].cells[hi]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_rtl(p)
        run = p.add_run(hdr)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        make_run_rtl(run)
        set_cell_shading(cell, '2C3E6B')
    
    for qi, q in enumerate(questions):
        q_cell = tbl.rows[qi+1].cells[0]
        a_cell = tbl.rows[qi+1].cells[1]
        q_cell.text = ''
        a_cell.text = ''
        
        # Find correct answer
        answer = '—'
        for line in q.split('\n'):
            clean_line = clean_inline(line)
            if '✔' in clean_line or '✓' in clean_line:
                # Extract the answer letter
                match = re.search(r'[abcd]\)\s*([^✔✓]+)', clean_line)
                if match:
                    answer = match.group(1).strip()
                else:
                    answer = clean_line.replace('✔','').replace('✓','').strip()
        
        # Question cell
        qp = q_cell.paragraphs[0]
        set_rtl(qp)
        qp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Remove the correct answer marker from displayed text
        display = q
        display = re.sub(r'\*\*.*?[✔✓].*?\*\*', '', display)
        display = display.replace('✔','').replace('✓','')
        display = display.strip()
        
        run = qp.add_run(display[:200])  # limit length
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        make_run_rtl(run)
        
        # Answer cell
        ap = a_cell.paragraphs[0]
        set_rtl(ap)
        ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = ap.add_run(answer)
        run2.font.size = Pt(9)
        run2.font.name = 'Calibri'
        run2.font.bold = True
        run2.font.color.rgb = RGBColor(0x1B, 0x7A, 0x2B)
        make_run_rtl(run2)
        
        if qi % 2 == 0:
            set_cell_shading(q_cell, 'F0F4FF')
            set_cell_shading(a_cell, 'F0F4FF')
    
    doc.add_paragraph()  # spacing after table


# ─── document builder ───

def create_semester_doc(lesson_files, output_name, semester_title, part_label):
    """Create a DOCX file for a batch of lessons."""
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # Set section RTL
    sectPr = section._sectPr
    # Add bidi to section
    rtl_sect = OxmlElement('w:bidi')
    rtl_sect.set(qn('w:val'), '1')
    sectPr.append(rtl_sect)
    
    # Remove docGrid that forces left-to-right
    docGrid = sectPr.find(qn('w:docGrid'))
    if docGrid is not None:
        sectPr.remove(docGrid)
    
    # Set text direction on section
    textDir = OxmlElement('w:textDirection')
    textDir.set(qn('w:val'), 'lrTb')
    sectPr.append(textDir)
    
    add_page_number(doc)
    
    # ── COVER PAGE ──
    for _ in range(6):
        doc.add_paragraph()
    
    new_para(doc, 'منهاج سيبليوس الشامل', bold=True, size=28, 
             color=(0x1A, 0x1A, 0x2E), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    new_para(doc, 'منهاج شامل لجميع الآلات الموسيقية', bold=True, size=18, 
             color=(0x2C, 0x3E, 0x6B), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    new_para(doc, 'Comprehensive Sibelius Curriculum — For All Instruments', 
             size=14, color=(0x66, 0x66, 0x66), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    new_para(doc, semester_title, bold=True, size=14, color=(0x34, 0x49, 0x7E), 
             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    new_para(doc, f'{len(lesson_files)} دروس', size=12, color=(0x88, 0x88, 0x88), 
             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    new_para(doc, 'إعداد: د. مصطفى كمال — Fujairah Academy of Fine Arts', size=12, 
             color=(0x34, 0x49, 0x7E), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    doc.add_page_break()
    
    # ── TABLE OF CONTENTS ──
    add_heading(doc, 'فهرس المحتويات', 1)
    doc.add_paragraph()
    
    for fpath in lesson_files:
        fname = os.path.basename(fpath)
        lesson_num = re.search(r'(\d+)', fname)
        num = lesson_num.group(1) if lesson_num else '?'
        
        with open(fpath, 'r', encoding='utf-8') as f:
            first_heading = '—'
            for line in f:
                if line.startswith('# '):
                    first_heading = line.strip('# \n')
                    break
        
        clean_title = clean_inline(first_heading)
        p = doc.add_paragraph()
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f'الدرس {num}')
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.name = 'Calibri'
        make_run_rtl(run)
        run2 = p.add_run(f' — {clean_title}')
        run2.font.size = Pt(11)
        run2.font.name = 'Calibri'
        run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        make_run_rtl(run2)
    
    doc.add_page_break()
    
    # ── LESSONS ──
    add_heading(doc, part_label, 1)
    doc.add_paragraph()
    
    for i, fpath in enumerate(lesson_files):
        if i > 0:
            doc.add_page_break()
        
        with open(fpath, 'r', encoding='utf-8') as f:
            content = f.read()
        
        process_lesson(doc, content)
    
    output_path = os.path.join(OUTPUT_DIR, output_name)
    doc.save(output_path)
    print(f"✅ {output_name}  ({os.path.getsize(output_path)//1024} KB)")
    return output_path


def main():
    s1 = sorted(glob.glob(os.path.join(SEMESTER1_DIR, "*.md")))
    s2 = sorted(glob.glob(os.path.join(SEMESTER2_DIR, "*.md")))
    
    print(f"Semester 1: {len(s1)} lessons")
    print(f"Semester 2: {len(s2)} lessons")
    
    # Split into 4 batches per semester (8 lessons each)
    s1_batches = [s1[i:i+8] for i in range(0, 32, 8)]
    s2_batches = [s2[i:i+8] for i in range(0, 32, 8)]
    
    s1_info = [
        ("Sibelius-Semester1-Part1-Lessons-1-8.docx", "الجزء الأول: الدروس 1–8 — الأساسيات", "الفصل الدراسي الأول — الجزء الأول"),
        ("Sibelius-Semester1-Part2-Lessons-9-16.docx", "الجزء الثاني: الدروس 9–16 — المهارات الأساسية", "الفصل الدراسي الأول — الجزء الثاني"),
        ("Sibelius-Semester1-Part3-Lessons-17-24.docx", "الجزء الثالث: الدروس 17–24 — التدوين والتنسيق", "الفصل الدراسي الأول — الجزء الثالث"),
        ("Sibelius-Semester1-Part4-Lessons-25-32.docx", "الجزء الرابع: الدروس 25–32 — متقدم الفصل الأول", "الفصل الدراسي الأول — الجزء الرابع"),
    ]
    
    s2_info = [
        ("Sibelius-Semester2-Part1-Lessons-33-40.docx", "الجزء الأول: الدروس 33–40 — تقنيات متوسطة", "الفصل الدراسي الثاني — الجزء الأول"),
        ("Sibelius-Semester2-Part2-Lessons-41-48.docx", "الجزء الثاني: الدروس 41–48 — مشاريع متقدمة", "الفصل الدراسي الثاني — الجزء الثاني"),
        ("Sibelius-Semester2-Part3-Lessons-49-56.docx", "الجزء الثالث: الدروس 49–56 — سير عمل احترافي", "الفصل الدراسي الثاني — الجزء الثالث"),
        ("Sibelius-Semester2-Part4-Lessons-57-64.docx", "الجزء الرابع: الدروس 57–64 — مشاريع التخرج والإتقان", "الفصل الدراسي الثاني — الجزء الرابع"),
    ]
    
    outputs = []
    
    print("\n📘 Semester 1...")
    for i, batch in enumerate(s1_batches):
        fn, title, part = s1_info[i]
        outputs.append(create_semester_doc(batch, fn, title, part))
    
    print("\n📗 Semester 2...")
    for i, batch in enumerate(s2_batches):
        fn, title, part = s2_info[i]
        outputs.append(create_semester_doc(batch, fn, title, part))
    
    total_kb = sum(os.path.getsize(p) // 1024 for p in outputs)
    print(f"\n{'='*50}")
    print(f"📦 {len(outputs)} files, {total_kb} KB total")
    print(f"📂 {OUTPUT_DIR}/")
    print(f"{'='*50}")

if __name__ == "__main__":
    main()
